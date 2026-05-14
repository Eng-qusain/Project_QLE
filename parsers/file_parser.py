"""
project_QLE/parsers/file_parser.py
─────────────────────────────
Unified parser for PDF, DOCX, XML, JPG, CSV.
Returns a ParsedFile regardless of input format.
"""

from __future__ import annotations

import io
import csv
import json
import logging
from pathlib import Path
from typing import Any, Dict, Optional

import pandas as pd

from project_QLE.core.models import FileType, ParsedFile

logger = logging.getLogger(__name__)


# ─────────────────────────────────────────────
#  Dispatch helper
# ─────────────────────────────────────────────

_EXTENSION_MAP: Dict[str, FileType] = {
    ".pdf"  : FileType.PDF,
    ".docx" : FileType.DOCX,
    ".doc"  : FileType.DOCX,
    ".xml"  : FileType.XML,
    ".jpg"  : FileType.JPG,
    ".jpeg" : FileType.JPG,
    ".png"  : FileType.JPG,
    ".csv"  : FileType.CSV,
    ".las"  : FileType.LAS,
    ".segy" : FileType.SEGY,
    ".sgy"  : FileType.SEGY,
}


def detect_file_type(path: Path) -> FileType:
    return _EXTENSION_MAP.get(path.suffix.lower(), FileType.UNKNOWN)


# ─────────────────────────────────────────────
#  Individual parsers
# ─────────────────────────────────────────────

def _parse_pdf(path: Path) -> ParsedFile:
    """Extract text and images from PDF using PyMuPDF (fitz)."""
    try:
        import fitz  # PyMuPDF
    except ImportError:
        raise ImportError("Install PyMuPDF: pip install PyMuPDF")

    doc = fitz.open(str(path))
    pages_text: list[str] = []
    images: list[bytes] = []
    meta: Dict[str, Any] = {
        "n_pages"  : doc.page_count,
        "title"    : doc.metadata.get("title", ""),
        "author"   : doc.metadata.get("author", ""),
        "subject"  : doc.metadata.get("subject", ""),
    }

    for page in doc:
        pages_text.append(page.get_text())
        for img_info in page.get_images(full=True):
            xref = img_info[0]
            base_image = doc.extract_image(xref)
            images.append(base_image["image"])

    doc.close()
    return ParsedFile(
        source_path=path,
        file_type=FileType.PDF,
        raw_text="\n".join(pages_text),
        metadata=meta,
        images=images,
    )


def _parse_docx(path: Path) -> ParsedFile:
    """Extract text and tables from DOCX."""
    try:
        from docx import Document
    except ImportError:
        raise ImportError("Install python-docx: pip install python-docx")

    doc = Document(str(path))
    paragraphs = [p.text for p in doc.paragraphs]

    tables_data: list[pd.DataFrame] = []
    for table in doc.tables:
        rows = [[cell.text for cell in row.cells] for row in table.rows]
        if rows:
            df = pd.DataFrame(rows[1:], columns=rows[0])
            tables_data.append(df)

    meta = {
        "n_paragraphs": len(paragraphs),
        "n_tables"    : len(tables_data),
        "core_props"  : {
            "author" : doc.core_properties.author,
            "title"  : doc.core_properties.title,
        },
    }

    return ParsedFile(
        source_path=path,
        file_type=FileType.DOCX,
        raw_text="\n".join(paragraphs),
        dataframe=tables_data[0] if tables_data else None,
        metadata=meta,
        extra={"all_tables": tables_data},
    )


def _parse_xml(path: Path) -> ParsedFile:
    """Parse XML, flatten to a DataFrame if structure is tabular."""
    try:
        from lxml import etree
    except ImportError:
        raise ImportError("Install lxml: pip install lxml")

    tree = etree.parse(str(path))
    root = tree.getroot()
    raw_text = etree.tostring(root, pretty_print=True, encoding="unicode")

    # Try flattening: collect all leaf elements
    records: list[Dict[str, str]] = []
    for child in root:
        record = {sub.tag: sub.text for sub in child}
        if record:
            records.append(record)

    df = pd.DataFrame(records) if records else None
    meta = {
        "root_tag" : root.tag,
        "n_children": len(root),
        "ns"       : root.nsmap,
    }
    return ParsedFile(
        source_path=path,
        file_type=FileType.XML,
        raw_text=raw_text,
        dataframe=df,
        metadata=meta,
    )


def _parse_image(path: Path) -> ParsedFile:
    """Read image, extract basic metadata, keep raw bytes."""
    try:
        from PIL import Image as PILImage
        import numpy as np
    except ImportError:
        raise ImportError("Install Pillow: pip install Pillow")

    img = PILImage.open(str(path))
    meta = {
        "format" : img.format,
        "mode"   : img.mode,
        "size"   : img.size,
    }
    buf = io.BytesIO()
    img.save(buf, format=img.format or "JPEG")
    raw_bytes = buf.getvalue()

    # Convert to numpy for stats
    import numpy as np
    arr = np.array(img)
    if arr.ndim == 3:
        meta["mean_rgb"] = arr.mean(axis=(0, 1)).tolist()
    img.close()

    return ParsedFile(
        source_path=path,
        file_type=FileType.JPG,
        images=[raw_bytes],
        metadata=meta,
    )


def _parse_csv(path: Path) -> ParsedFile:
    """Read CSV into a DataFrame, infer numeric columns."""
    df = pd.read_csv(str(path), low_memory=False)
    df = df.apply(pd.to_numeric, errors="ignore")

    meta = {
        "n_rows"   : len(df),
        "n_cols"   : len(df.columns),
        "columns"  : list(df.columns),
        "dtypes"   : {c: str(t) for c, t in df.dtypes.items()},
    }
    return ParsedFile(
        source_path=path,
        file_type=FileType.CSV,
        dataframe=df,
        metadata=meta,
    )


# ─────────────────────────────────────────────
#  Public API
# ─────────────────────────────────────────────

_PARSERS = {
    FileType.PDF  : _parse_pdf,
    FileType.DOCX : _parse_docx,
    FileType.XML  : _parse_xml,
    FileType.JPG  : _parse_image,
    FileType.CSV  : _parse_csv,
}


def parse_file(path: str | Path) -> ParsedFile:
    """
    Parse any supported file and return a unified ParsedFile object.

    Parameters
    ----------
    path : str or Path
        Path to the file on disk.

    Returns
    -------
    ParsedFile
    """
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {path}")

    file_type = detect_file_type(path)
    parser = _PARSERS.get(file_type)

    if parser is None:
        logger.warning("No dedicated parser for %s (%s). Returning stub.", path.name, file_type)
        return ParsedFile(source_path=path, file_type=file_type)

    logger.info("Parsing %s as %s …", path.name, file_type.value.upper())
    return parser(path)